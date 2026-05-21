import json
import shutil
from pathlib import Path
from docx import Document


def _get_paragraph(doc, segment: dict):
    """Navigate to the exact paragraph for a segment using its container metadata."""
    ct = segment["container_type"]
    ci = segment.get("container_index")
    lpi = segment.get("container_paragraph_index")

    if ct == "paragraph":
        if lpi is not None and lpi < len(doc.paragraphs):
            return doc.paragraphs[lpi]
    elif ct == "table_cell":
        tr = segment.get("table_row")
        tc = segment.get("table_col")
        if ci is not None and ci < len(doc.tables):
            table = doc.tables[ci]
            if tr is not None and tc is not None:
                if tr < len(table.rows) and tc < len(table.rows[tr].cells):
                    cell = table.rows[tr].cells[tc]
                    if lpi is not None and lpi < len(cell.paragraphs):
                        return cell.paragraphs[lpi]
    elif ct == "header":
        si = segment.get("section_index")
        if si is not None and si < len(doc.sections):
            header = doc.sections[si].header
            if header and lpi is not None and lpi < len(header.paragraphs):
                return header.paragraphs[lpi]
    elif ct == "footer":
        si = segment.get("section_index")
        if si is not None and si < len(doc.sections):
            footer = doc.sections[si].footer
            if footer and lpi is not None and lpi < len(footer.paragraphs):
                return footer.paragraphs[lpi]
    elif ct == "header_table_cell":
        si = segment.get("section_index")
        tr = segment.get("table_row")
        tc = segment.get("table_col")
        if si is not None and si < len(doc.sections):
            header = doc.sections[si].header
            if header and ci is not None and ci < len(header.tables):
                table = header.tables[ci]
                if tr is not None and tc is not None:
                    if tr < len(table.rows) and tc < len(table.rows[tr].cells):
                        cell = table.rows[tr].cells[tc]
                        if lpi is not None and lpi < len(cell.paragraphs):
                            return cell.paragraphs[lpi]
    elif ct == "footer_table_cell":
        si = segment.get("section_index")
        tr = segment.get("table_row")
        tc = segment.get("table_col")
        if si is not None and si < len(doc.sections):
            footer = doc.sections[si].footer
            if footer and ci is not None and ci < len(footer.tables):
                table = footer.tables[ci]
                if tr is not None and tc is not None:
                    if tr < len(table.rows) and tc < len(table.rows[tr].cells):
                        cell = table.rows[tr].cells[tc]
                        if lpi is not None and lpi < len(cell.paragraphs):
                            return cell.paragraphs[lpi]
    return None


def generate_translated_docx(
    project_id: int,
    target_lang: str,
    db,
    original_file_path: str,
    output_file_path: str,
) -> str:
    # Copy original to preserve all formatting, styles, images, etc.
    shutil.copy2(original_file_path, output_file_path)
    doc = Document(output_file_path)

    rows = db.execute(
        """SELECT s.* FROM segments s
           WHERE s.project_id = ? AND s.is_translated = 1 AND s.ignored = 0
           AND s.key_id IN (
             SELECT tv.key_id FROM translation_values tv WHERE tv.target_lang = ?
           )
           ORDER BY s.sequence""",
        (project_id, target_lang),
    ).fetchall()

    replaced = 0
    for row in rows:
        seg = dict(row)
        translated = (seg.get("translated_text") or "").strip()
        if not translated:
            continue
        fmt_json = seg.get("formatting_json") or "{}"
        if isinstance(fmt_json, str):
            fmt_json = json.loads(fmt_json)
        if fmt_json.get("contains_image"):
            continue

        paragraph = _get_paragraph(doc, seg)
        if paragraph is None:
            continue
        run_idx = seg["run_index"]
        if run_idx < len(paragraph.runs):
            paragraph.runs[run_idx].text = translated
            replaced += 1

    doc.save(output_file_path)
    return output_file_path
