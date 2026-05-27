import json
from docx import Document
from typing import List

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def extract_run_formatting(run) -> dict:
    font = run.font
    fmt = {}
    if font.name:
        fmt["font_name"] = font.name
    if font.size:
        fmt["font_size_pt"] = font.size.pt
    if font.bold is not None:
        fmt["bold"] = font.bold
    if font.italic is not None:
        fmt["italic"] = font.italic
    if font.underline is not None:
        fmt["underline"] = font.underline
    try:
        if font.color and font.color.rgb:
            fmt["color_rgb"] = str(font.color.rgb)
    except Exception:
        pass
    if font.superscript is not None:
        fmt["superscript"] = font.superscript
    if font.subscript is not None:
        fmt["subscript"] = font.subscript
    if font.strike is not None:
        fmt["strike"] = font.strike
    # Check for images
    drawings = run._element.findall(f"{{{WML_NS}}}drawing")
    if drawings:
        fmt["contains_image"] = True
    # Check for hyperlink
    parent = run._element.getparent()
    if parent is not None and parent.tag.endswith("}hyperlink"):
        fmt["is_hyperlink"] = True
    return fmt


def extract_paragraph_formatting(paragraph) -> dict:
    pf = paragraph.paragraph_format
    fmt = {}
    if pf.alignment is not None:
        fmt["alignment"] = str(pf.alignment)
    try:
        if pf.first_line_indent is not None:
            fmt["first_line_indent_pt"] = pf.first_line_indent.pt
        if pf.left_indent is not None:
            fmt["left_indent_pt"] = pf.left_indent.pt
    except Exception:
        pass
    try:
        if pf.space_before is not None:
            fmt["space_before_pt"] = pf.space_before.pt
        if pf.space_after is not None:
            fmt["space_after_pt"] = pf.space_after.pt
    except Exception:
        pass
    if pf.line_spacing is not None:
        try:
            fmt["line_spacing"] = float(pf.line_spacing)
        except Exception:
            fmt["line_spacing"] = str(pf.line_spacing)
    # Numbering
    pPr = paragraph._element.find(f"{{{WML_NS}}}pPr")
    if pPr is not None:
        numPr = pPr.find(f"{{{WML_NS}}}numPr")
        if numPr is not None:
            numId_el = numPr.find(f"{{{WML_NS}}}numId")
            level_el = numPr.find(f"{{{WML_NS}}}ilvl")
            if numId_el is not None:
                val = numId_el.get(f"{{{WML_NS}}}val")
                if val is not None:
                    fmt["numId"] = int(val)
            if level_el is not None:
                val = level_el.get(f"{{{WML_NS}}}val")
                if val is not None:
                    fmt["level"] = int(val)
    return fmt


def _iter_paragraph_runs(paragraph, para_idx, container_type, container_index,
                          table_row, table_col, local_para_idx, section_index=None):
    para_fmt = extract_paragraph_formatting(paragraph)
    para_fmt_json = json.dumps(para_fmt, ensure_ascii=False)

    # Collect all valid runs (only skip whitespace-only)
    raw = []
    for run_idx, run in enumerate(paragraph.runs):
        text = run.text
        if not text.strip():
            continue
        run_fmt = extract_run_formatting(run)
        run_fmt_json = json.dumps(run_fmt, ensure_ascii=False, sort_keys=True)
        raw.append({
            "run_index": run_idx,
            "text": text,
            "fmt_json": run_fmt_json,
        })

    # Merge consecutive runs with identical formatting
    merged = []
    for r in raw:
        if merged and merged[-1]["fmt_json"] == r["fmt_json"]:
            merged[-1]["text"] += r["text"]
            merged[-1]["run_count"] += 1
        else:
            merged.append({
                "run_index": r["run_index"],
                "text": r["text"],
                "fmt_json": r["fmt_json"],
                "run_count": 1,
            })

    segments = []
    for m in merged:
        segments.append({
            "paragraph_index": para_idx,
            "container_paragraph_index": local_para_idx,
            "run_index": m["run_index"],
            "run_count": m["run_count"],
            "source_text": m["text"],
            "formatting_json": m["fmt_json"],
            "paragraph_formatting_json": para_fmt_json,
            "container_type": container_type,
            "container_index": container_index,
            "table_row": table_row,
            "table_col": table_col,
            "section_index": section_index,
        })
    return segments


def parse_docx(file_path: str) -> List[dict]:
    doc = Document(file_path)
    segments = []
    seq = 0
    para_idx = 0

    # --- Body paragraphs ---
    for local_pi, paragraph in enumerate(doc.paragraphs):
        runs = _iter_paragraph_runs(
            paragraph, para_idx, "paragraph", None, None, None, local_pi
        )
        for r in runs:
            r["sequence"] = seq
            segments.append(r)
            seq += 1
        para_idx += 1

    # --- SDT (Structured Document Tags — TOC, content controls) ---
    # python-docx does not return paragraphs inside SDT via doc.paragraphs
    from lxml import etree
    body = doc.element.body
    sdt_containers = body.findall(f'.//{{{WML_NS}}}sdt/{{{WML_NS}}}sdtContent')
    for sdt_content in sdt_containers:
        sdt_paras = sdt_content.findall(f'{{{WML_NS}}}p')
        for local_pi, p_el in enumerate(sdt_paras):
            # Build a paragraph-like object from the XML element
            from docx.text.paragraph import Paragraph
            paragraph = Paragraph(p_el, None)
            runs = _iter_paragraph_runs(
                paragraph, para_idx, "paragraph", None, None, None, local_pi
            )
            for r in runs:
                r["sequence"] = seq
                segments.append(r)
                seq += 1
            para_idx += 1
        # Also handle tables inside SDT
        sdt_tables = sdt_content.findall(f'{{{WML_NS}}}tbl')
        for table_idx, tbl_el in enumerate(sdt_tables):
            from docx.table import Table
            table = Table(tbl_el, None)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for local_pi, paragraph in enumerate(cell.paragraphs):
                        runs = _iter_paragraph_runs(
                            paragraph, para_idx, "table_cell", table_idx,
                            row_idx, col_idx, local_pi
                        )
                        for r in runs:
                            r["sequence"] = seq
                            segments.append(r)
                            seq += 1
                        para_idx += 1

    # --- Tables ---
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for local_pi, paragraph in enumerate(cell.paragraphs):
                    runs = _iter_paragraph_runs(
                        paragraph, para_idx, "table_cell", table_idx,
                        row_idx, col_idx, local_pi
                    )
                    for r in runs:
                        r["sequence"] = seq
                        segments.append(r)
                        seq += 1
                    para_idx += 1

    # --- Headers & Footers ---
    for section_idx, section in enumerate(doc.sections):
        if section.header:
            for local_pi, paragraph in enumerate(section.header.paragraphs):
                runs = _iter_paragraph_runs(
                    paragraph, para_idx, "header", section_idx,
                    None, None, local_pi, section_idx
                )
                for r in runs:
                    r["sequence"] = seq
                    segments.append(r)
                    seq += 1
                para_idx += 1
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for local_pi, paragraph in enumerate(cell.paragraphs):
                            runs = _iter_paragraph_runs(
                                paragraph, para_idx, "header_table_cell", table_idx,
                                row_idx, col_idx, local_pi, section_idx
                            )
                            for r in runs:
                                r["sequence"] = seq
                                segments.append(r)
                                seq += 1
                            para_idx += 1

        if section.footer:
            for local_pi, paragraph in enumerate(section.footer.paragraphs):
                runs = _iter_paragraph_runs(
                    paragraph, para_idx, "footer", section_idx,
                    None, None, local_pi, section_idx
                )
                for r in runs:
                    r["sequence"] = seq
                    segments.append(r)
                    seq += 1
                para_idx += 1
            for table_idx, table in enumerate(section.footer.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for local_pi, paragraph in enumerate(cell.paragraphs):
                            runs = _iter_paragraph_runs(
                                paragraph, para_idx, "footer_table_cell", table_idx,
                                row_idx, col_idx, local_pi, section_idx
                            )
                            for r in runs:
                                r["sequence"] = seq
                                segments.append(r)
                                seq += 1
                            para_idx += 1

    return segments
